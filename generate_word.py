from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('📦 Catálogo de Core Assets — Fábrica de Software Académico', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Proyecto: Fábrica de Software Académico')
    doc.add_paragraph('Equipo: Jostin Quilca, Leonel Arellano, Alex Luna')
    doc.add_paragraph('Versión del catálogo: 1.2 (Smart Factory)')
    doc.add_paragraph('Fecha: Junio 2026')
    doc.add_paragraph('Línea de Producto: Sistemas Web Académicos con Autenticación y Auditoría')
    
    doc.add_heading('📋 Índice de Core Assets', level=1)
    
    # Index Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'ID'
    hdr_cells[2].text = 'Nombre'
    hdr_cells[3].text = 'Capa'
    hdr_cells[4].text = 'Reutilización'
    
    index_data = [
        ('1', 'CA-001', 'Sistema de Diseño (Design System)', 'Frontend', '100%'),
        ('2', 'CA-002', 'Modelo de Usuario (Frontend)', 'Frontend', '100%'),
        ('3', 'CA-003', 'Servicio de Autenticación', 'Frontend', '100%'),
        ('4', 'CA-004', 'Guarda de Autenticación', 'Frontend', '100%'),
        ('5', 'CA-005', 'Guarda de Roles', 'Frontend', '100%'),
        ('6', 'CA-006', 'Componente de Login', 'Frontend', 'Parcial'),
        ('7', 'CA-007', 'Componente de Registro', 'Frontend', 'Parcial'),
        ('8', 'CA-008', 'Componente Dashboard', 'Frontend', 'Parcial'),
        ('9', 'CA-009', 'Esquema GraphQL Base', 'Backend', '100%'),
        ('10', 'CA-010', 'Resolvers GraphQL', 'Backend', 'Parcial'),
        ('11', 'CA-011', 'Middleware de Autenticación JWT', 'Backend', '100%'),
        ('12', 'CA-012', 'Modelo de Auditoría', 'Backend', '100%'),
        ('13', 'CA-013', 'Configuración de Base de Datos', 'Backend', '100%'),
        ('14', 'CA-014', 'Configuración de la Fábrica', 'Herramienta', '100%'),
        ('15', 'CA-015', 'CLI de Ensamblaje Inteligente', 'Herramienta', '100%'),
    ]
    
    for item in index_data:
        row_cells = table.add_row().cells
        for i in range(5):
            row_cells[i].text = item[i]

    assets = [
        {
            "id": "CA-001", "name": "Sistema de Diseño (Design System)", "tech": "SCSS, CSS3, Google Fonts (Inter)", "file": "frontend/src/styles.scss",
            "desc": "Sistema de diseño visual reutilizable que define la identidad gráfica de todas las aplicaciones producidas por la fábrica. Incluye variables de color, tipografía, componentes UI base y efectos visuales premium (Glassmorphism).",
            "func": ["Paleta de colores configurable mediante variables CSS/SCSS", "Tipografía base Inter importada desde Google Fonts", "Componentes visuales: .glass-panel, .btn, .btn-primary, .form-control, .form-group", "Alertas: .alert-error, .alert-success", "Efectos: transiciones suaves, hover effects, gradientes", "Responsive design base"],
            "var": [("Colores primarios", "Configuración", "Cambiar variables $primary, $secondary para adaptar a otra marca"), ("Tipografía", "Configuración", "Sustituir Inter por otra fuente en el @import"), ("Efectos glassmorphism", "Opcional", "Activar/desactivar el efecto cristal según el estilo del proyecto")],
            "integ": "1. Copiar styles.scss al nuevo proyecto Angular\n2. Ajustar las variables de color según la identidad del nuevo proyecto\n3. Importar en angular.json dentro de styles[]",
            "deps": "Google Fonts (Inter) — conexión a internet o descarga local",
            "val": ["Los componentes .glass-panel, .btn, .form-control se renderizan correctamente", "La tipografía Inter carga sin errores en consola", "Los colores se aplican consistentemente en todas las páginas"]
        },
        {
            "id": "CA-002", "name": "Modelo de Usuario (Frontend)", "tech": "TypeScript, Angular", "file": "frontend/src/app/models/user.model.ts",
            "desc": "Define las interfaces y enumeraciones TypeScript que modelan la estructura del usuario en el frontend. Proporciona tipado fuerte para toda la capa de presentación.",
            "func": ["Enum Role con valores: ADMIN, DOCENTE, ESTUDIANTE", "Interface Usuario con campos: id, nombre, email, rol, createdAt", "Interface AuthPayload con campos: token, usuario"],
            "var": [("Roles", "Extensión", "Agregar nuevos roles al enum (ej. COORDINADOR, TUTOR)"), ("Campos de Usuario", "Extensión", "Agregar propiedades opcionales (ej. avatar, telefono)")],
            "integ": "1. Copiar user.model.ts a src/app/models/\n2. Ajustar los roles del enum según el nuevo proyecto\n3. Importar donde se necesite.",
            "deps": "Ninguna (solo TypeScript nativo)",
            "val": ["El enum Role contiene todos los roles requeridos por el nuevo proyecto", "La interface Usuario coincide con la estructura del backend"]
        },
        {
            "id": "CA-003", "name": "Servicio de Autenticación", "tech": "Angular, RxJS, HttpClient", "file": "frontend/src/app/services/auth.service.ts",
            "desc": "Servicio Angular inyectable que encapsula toda la lógica de autenticación: login, registro, logout, persistencia de sesión y estado reactivo del usuario actual.",
            "func": ["login(email, password) — Envía mutación GraphQL y almacena token JWT", "registro(nombre, email, password, rol) — Registra usuario nuevo", "logout() — Limpia sesión del localStorage", "isAuthenticated() — Verifica si hay sesión activa", "getUserRole() — Retorna el rol del usuario actual", "currentUser — Signal reactivo de Angular para el estado del usuario", "Manejo correcto de errores GraphQL (HTTP 200 con errors[])"],
            "var": [("URL de la API", "Configuración", "Cambiar environment.apiUrl al endpoint del nuevo backend"), ("Campos del registro", "Extensión", "Agregar campos adicionales a la mutación de registro"), ("Tipo de API", "Reemplazo", "Cambiar de GraphQL a REST modificando los métodos HTTP")],
            "integ": "1. Copiar auth.service.ts a src/app/services/\n2. Configurar environment.ts con la URL de la API del nuevo proyecto\n3. Asegurarse de importar HttpClientModule en la configuración de la app\n4. Ajustar las mutaciones GraphQL si el schema del backend es diferente",
            "deps": "CA-002 (Modelo de Usuario), @angular/common/http (HttpClient), rxjs",
            "val": ["Login exitoso almacena token en localStorage", "Login fallido muestra mensaje de error sin quedarse cargando", "isAuthenticated() retorna true tras login exitoso", "logout() limpia completamente la sesión"]
        },
        {
            "id": "CA-004", "name": "Guarda de Autenticación (AuthGuard)", "tech": "Angular Router", "file": "frontend/src/app/guards/auth.guard.ts",
            "desc": "Guarda de ruta funcional que protege rutas privadas. Verifica si el usuario tiene una sesión activa antes de permitir el acceso. Si no está autenticado, redirige automáticamente a /login.",
            "func": ["Verificación de sesión mediante AuthService.isAuthenticated()", "Redirección automática a /login si no hay sesión", "Compatible con canActivate del Angular Router"],
            "var": [("Ruta de redirección", "Configuración", "Cambiar /login por otra ruta si el proyecto lo requiere")],
            "integ": "1. Copiar auth.guard.ts a src/app/guards/\n2. Aplicar en app.routes.ts configurando canActivate: [authGuard]",
            "deps": "CA-003 (Servicio de Autenticación)",
            "val": ["Usuarios no autenticados son redirigidos a /login", "Usuarios autenticados acceden normalmente a rutas protegidas"]
        },
        {
            "id": "CA-005", "name": "Guarda de Roles (RoleGuard)", "tech": "Angular Router", "file": "frontend/src/app/guards/role.guard.ts",
            "desc": "Guarda de ruta que restringe el acceso a páginas específicas basándose en el rol del usuario autenticado. Permite configurar qué roles tienen permiso para acceder a cada ruta.",
            "func": ["Lectura del rol del usuario desde AuthService.getUserRole()", "Comparación contra roles permitidos definidos en route.data", "Redirección a /dashboard si el rol no tiene permiso"],
            "var": [("Roles permitidos", "Configuración", "Definir roles por ruta en route.data['roles']"), ("Ruta de rechazo", "Configuración", "Cambiar /dashboard por otra ruta de fallback")],
            "integ": "1. Copiar role.guard.ts a src/app/guards/\n2. Configurar en las rutas con canActivate: [authGuard, roleGuard] y data: { roles: ['ADMIN'] }",
            "deps": "CA-003 (Servicio de Autenticación), CA-004 (Guarda de Autenticación)",
            "val": ["Un ESTUDIANTE no puede acceder a rutas marcadas solo para ADMIN", "Un ADMIN accede correctamente a todas las rutas que le corresponden"]
        },
        {
            "id": "CA-006", "name": "Componente de Login", "tech": "Angular, Reactive Forms, ChangeDetectorRef", "file": "frontend/src/app/pages/login/ (.ts, .html, .scss)",
            "desc": "Componente standalone de Angular que presenta la interfaz de inicio de sesión. Incluye validaciones reactivas, toggle de visibilidad de contraseña, manejo de estados de carga y visualización de errores del servidor.",
            "func": ["Formulario reactivo con validaciones", "Botón de mostrar/ocultar contraseña", "Estado de carga con spinner visual", "Alerta de error que muestra mensajes del backend en tiempo real", "Enlace a recuperación de contraseña y registro", "ChangeDetectorRef para forzar repintado tras errores asíncronos"],
            "var": [("Diseño visual", "Configuración", "Modificar el .scss para adaptar a la marca del nuevo proyecto"), ("Campos del formulario", "Extensión", "Agregar campos como recordar sesión o captcha"), ("Ruta post-login", "Configuración", "Cambiar /dashboard por la ruta destino del nuevo proyecto")],
            "integ": "1. Copiar la carpeta pages/login/ completa\n2. Ajustar colores en el .scss para coincidir con CA-001\n3. Verificar que la ruta post-login en onSubmit() apunte al destino correcto\n4. Registrar la ruta en app.routes.ts",
            "deps": "CA-001 (Sistema de Diseño), CA-003 (Servicio de Autenticación)",
            "val": ["Login exitoso redirige al dashboard", "Credenciales incorrectas muestran alerta roja inmediatamente", "El botón se desactiva mientras el formulario es inválido"]
        },
        {
            "id": "CA-007", "name": "Componente de Registro", "tech": "Angular, Reactive Forms, Custom Validators", "file": "frontend/src/app/pages/registro/ (.ts, .html, .scss)",
            "desc": "Componente standalone de Angular para el registro de nuevos usuarios. Incluye validadores personalizados para correo institucional y contraseña segura, con retroalimentación visual en tiempo real.",
            "func": ["Formulario reactivo con campos completos", "Validador personalizado institucionalEmailValidator (exige .edu o .edu.ec)", "Validador personalizado strongPasswordValidator (mayúsculas, números, símbolos)", "Validación de coincidencia de contraseñas", "Selector de rol (ADMIN, DOCENTE, ESTUDIANTE)"],
            "var": [("Dominio del email", "Configuración", "Cambiar .edu.ec por otro dominio institucional en el validador"), ("Reglas de contraseña", "Configuración", "Ajustar requisitos mínimos en strongPasswordValidator"), ("Roles disponibles", "Configuración", "Modificar las opciones del selector de roles")],
            "integ": "1. Copiar la carpeta pages/registro/ completa\n2. Modificar el validador de email para el dominio institucional requerido\n3. Ajustar los roles disponibles en el HTML",
            "deps": "CA-001 (Sistema de Diseño), CA-002 (Modelo de Usuario), CA-003 (Servicio de Autenticación)",
            "val": ["Correos fuera del dominio institucional son rechazados visualmente", "Contraseñas débiles muestran alerta con los requisitos faltantes", "Registro exitoso redirige y crea sesión automáticamente"]
        },
        {
            "id": "CA-008", "name": "Componente Dashboard", "tech": "Angular, Signals, *ngIf", "file": "frontend/src/app/pages/dashboard/ (.ts, .html, .scss)",
            "desc": "Panel de control dinámico que adapta su contenido visual según el rol del usuario autenticado. Utiliza Angular Signals para leer el estado del usuario en tiempo real.",
            "func": ["Lectura del usuario actual desde AuthService.currentUser (Signal)", "Renderizado condicional de módulos por rol (*ngIf)", "Botón de logout funcional", "Tarjetas de módulos diferenciadas por rol"],
            "var": [("Módulos por rol", "Extensión", "Agregar/quitar tarjetas de módulos según los requerimientos del nuevo proyecto"), ("Layout", "Configuración", "Cambiar de grid a sidebar, tabs, etc.")],
            "integ": "1. Copiar la carpeta pages/dashboard/ completa\n2. Modificar las secciones *ngIf del HTML para reflejar los módulos del nuevo proyecto",
            "deps": "CA-001 (Sistema de Diseño), CA-003 (Servicio de Autenticación)",
            "val": ["Un ADMIN ve solo módulos de administración", "Un ESTUDIANTE ve solo módulos de estudiante", "El botón de logout cierra sesión y redirige al login"]
        },
        {
            "id": "CA-009", "name": "Esquema GraphQL Base", "tech": "GraphQL, Apollo Server", "file": "backend/src/schema/typeDefs.js",
            "desc": "Definición del esquema GraphQL que establece los tipos de datos, queries y mutaciones disponibles en la API. Sirve como contrato entre el frontend y el backend.",
            "func": ["Tipos: Usuario, Role, AuthPayload, Auditoria", "Queries: usuarios, usuario(id), roles, me, auditoria, auditoriaByUsuario, auditoriaByAccion", "Mutaciones: registro, login, logout"],
            "var": [("Tipos de datos", "Extensión", "Agregar nuevos tipos (ej. Curso, Nota, Horario)"), ("Queries/Mutaciones", "Extensión", "Agregar nuevas consultas y operaciones CRUD para nuevas entidades")],
            "integ": "1. Copiar typeDefs.js como base del nuevo proyecto\n2. Mantener los tipos base intactos\n3. Agregar nuevos tipos y operaciones según el dominio del proyecto",
            "deps": "Apollo Server (@apollo/server)",
            "val": ["El servidor arranca sin errores de esquema", "Todas las queries y mutaciones son accesibles desde el playground de Apollo"]
        },
        {
            "id": "CA-010", "name": "Resolvers GraphQL", "tech": "Node.js, bcryptjs, jsonwebtoken, pg", "file": "backend/src/resolvers/index.js",
            "desc": "Lógica de negocio del backend que implementa las operaciones definidas en el esquema GraphQL. Incluye autenticación segura con hash de contraseñas, generación de JWT y registro automático de eventos de auditoría.",
            "func": ["Resolvers de campo: Usuario.rol, Usuario.createdAt", "Resolvers de auditoría: mapeo de campos", "Mutación registro: hash bcrypt + creación de usuario + auditoría", "Mutación login: verificación de credenciales + JWT + auditoría de éxito/fallo", "Mutación logout: invalidación de sesión + auditoría"],
            "var": [("Lógica de registro", "Extensión", "Agregar validaciones adicionales"), ("Expiración del token", "Configuración", "Cambiar '24h' por otro valor"), ("Nuevos resolvers", "Extensión", "Agregar resolvers para nuevas entidades del proyecto")],
            "integ": "1. Copiar resolvers/index.js como base\n2. Mantener la sección de autenticación intacta\n3. Agregar nuevos resolvers para las entidades específicas del proyecto",
            "deps": "CA-009 (Esquema GraphQL), CA-012 (Modelo de Auditoría), CA-013 (Configuración de BD)",
            "val": ["Login con credenciales correctas retorna un token JWT válido", "Login con credenciales incorrectas registra evento LOGIN_FALLIDO", "Registro crea usuario con contraseña encriptada"]
        },
        {
            "id": "CA-011", "name": "Middleware de Autenticación JWT", "tech": "Node.js, jsonwebtoken", "file": "backend/src/middleware/auth.js",
            "desc": "Middleware que intercepta cada petición HTTP entrante, extrae el token JWT del header Authorization, lo verifica y decodifica los datos del usuario para inyectarlos en el contexto de GraphQL.",
            "func": ["Extracción del token del header Authorization", "Verificación y decodificación del JWT usando JWT_SECRET", "Retorno de null si el token es inválido o no existe"],
            "var": [("Secret del JWT", "Configuración", "Cambiar JWT_SECRET en el .env"), ("Formato del token", "Configuración", "Soportar Bearer <token> o token directo")],
            "integ": "1. Copiar middleware/auth.js al nuevo proyecto\n2. Configurar JWT_SECRET en las variables de entorno\n3. Integrar en el contexto de Apollo Server",
            "deps": "jsonwebtoken, Variable de entorno JWT_SECRET",
            "val": ["Peticiones con token válido inyectan user en el contexto", "Peticiones sin token retornan user: null sin error"]
        },
        {
            "id": "CA-012", "name": "Modelo de Auditoría", "tech": "Node.js, pg (PostgreSQL)", "file": "backend/src/models/Auditoria.js",
            "desc": "Modelo de datos que gestiona el registro y consulta de eventos de auditoría del sistema. Permite rastrear quién hizo qué, cuándo y desde dónde.",
            "func": ["registrar() — Inserta un evento", "findAll(), findByUsuario(), findByAccion() — Consultas con paginación y filtros", "Tipos de acción soportados: LOGIN, LOGOUT, LOGIN_FALLIDO, REGISTRO, ACCESO_RUTA"],
            "var": [("Tipos de acción", "Extensión", "Agregar nuevas acciones (ej. CREAR_CURSO)"), ("Campos adicionales", "Extensión", "Agregar user_agent, geolocation, etc.")],
            "integ": "1. Copiar models/Auditoria.js al nuevo proyecto\n2. Crear la tabla auditoria en la base de datos\n3. Importar y usar en los resolvers",
            "deps": "CA-013 (Configuración de BD), Tabla auditoria creada en PostgreSQL",
            "val": ["Cada login exitoso genera un registro con acción LOGIN", "Cada intento fallido genera un registro con acción LOGIN_FALLIDO"]
        },
        {
            "id": "CA-013", "name": "Configuración de Base de Datos", "tech": "Node.js, pg (PostgreSQL), dotenv", "file": "backend/src/config/database.js, backend/.env",
            "desc": "Módulo de configuración que establece la conexión con la base de datos PostgreSQL en la nube. Utiliza variables de entorno para mantener las credenciales seguras y soporta conexiones SSL.",
            "func": ["Conexión a PostgreSQL mediante pg.Client", "Soporte SSL para bases de datos en la nube", "Función connectDB() con verificación de conexión"],
            "var": [("Proveedor de BD", "Configuración", "Cambiar credenciales en .env para Render, Supabase, AWS, etc."), ("SSL", "Configuración", "Activar/desactivar según el proveedor")],
            "integ": "1. Copiar config/database.js al nuevo proyecto\n2. Crear archivo .env con las credenciales\n3. Llamar connectDB() al iniciar el servidor",
            "deps": "pg (node-postgres), dotenv",
            "val": ["connectDB() imprime la fecha actual del servidor PostgreSQL", "El archivo .env está en .gitignore", "La conexión funciona con SSL habilitado"]
        },
        {
            "id": "CA-014", "name": "Configuración de la Fábrica (Feature Toggles)", "tech": "JSON", "file": "factory-config.json",
            "desc": "Punto de control de variabilidad de la línea de producción. Enumera explícitamente los 13 Core Assets permitiendo habilitar o deshabilitar funcionalidades opcionales mediante valores booleanos.",
            "func": ["Listado completo de los 13 Core Assets", "Control de funcionalidades opcionales (ej. CA-012_ModeloAuditoria: false)", "Controla configuraciones de entorno iniciales (DB, Puertos)"],
            "var": [("Valores booleanos por Asset", "Configuración", "Apagar o encender funcionalidades opcionales en el nuevo proyecto")],
            "integ": "1. Mantener en la raíz de la Fábrica de Software\n2. Cambiar a 'false' los módulos no requeridos antes de ensamblar el proyecto",
            "deps": "Ninguna",
            "val": ["Formato JSON válido", "Contiene el bloque core_assets con las 13 llaves"]
        },
        {
            "id": "CA-015", "name": "CLI de Ensamblaje Inteligente (Smart Scaffolding)", "tech": "Node.js, FileSystem (fs)", "file": "crear_nueva_app.js",
            "desc": "Script de línea de comandos avanzado que actúa como la 'cadena de montaje'. Clona los Core Assets y realiza una poda inteligente del código fuente basándose en la configuración de CA-014.",
            "func": ["Validación de seguridad: advierte y evita la eliminación de assets obligatorios (Commonalities)", "Poda Inteligente: Elimina modelos, rutas y referencias GraphQL de los assets apagados (Variabilities) mediante AST-level Regex", "Inyección automática de entorno base (.env)"],
            "var": [("Reglas de Exclusión (Poda)", "Extensión", "Añadir lógica de regex para limpiar código de nuevos assets opcionales creados en el futuro")],
            "integ": "1. Ejecutar en la terminal desde la raíz usando: node crear_nueva_app.js NombreDelProyecto",
            "deps": "CA-014 (Configuración de la Fábrica)",
            "val": ["Mantiene assets obligatorios aunque se configuren en false", "Elimina de forma quirúrgica rastros de código de assets marcados en false"]
        }
    ]

    for asset in assets:
        doc.add_page_break()
        doc.add_heading(f"{asset['id']} — {asset['name']}", level=2)
        
        doc.add_heading('Identificación', level=3)
        t_id = doc.add_table(rows=6, cols=2)
        t_id.style = 'Table Grid'
        data_id = [
            ("ID", asset['id']),
            ("Tecnologías", asset['tech']),
            ("Archivo(s)", asset['file']),
            ("Estado", "Estable"),
            ("Responsable", "Equipo Core"),
            ("Versión", "1.0")
        ]
        for i, row in enumerate(t_id.rows):
            row.cells[0].text = data_id[i][0]
            row.cells[1].text = data_id[i][1]
            
        doc.add_heading('Descripción', level=3)
        doc.add_paragraph(asset['desc'])
        
        doc.add_heading('Funcionalidades', level=3)
        for f in asset['func']:
            doc.add_paragraph(f, style='List Bullet')
            
        doc.add_heading('Puntos de Variabilidad', level=3)
        t_var = doc.add_table(rows=1, cols=3)
        t_var.style = 'Table Grid'
        hdr = t_var.rows[0].cells
        hdr[0].text = 'Punto de Variación'
        hdr[1].text = 'Tipo'
        hdr[2].text = 'Descripción'
        for v in asset['var']:
            r = t_var.add_row().cells
            r[0].text = v[0]
            r[1].text = v[1]
            r[2].text = v[2]
            
        doc.add_heading('Guía de Integración', level=3)
        doc.add_paragraph(asset['integ'])
        
        doc.add_heading('Dependencias', level=3)
        doc.add_paragraph(asset['deps'])
        
        doc.add_heading('Criterios de Validación', level=3)
        for v in asset['val']:
            doc.add_paragraph(f"[ ] {v}", style='List Bullet')

    doc.save('Documentacion_Core_Assets_v3.docx')
    print("Word document updated successfully with detailed CA-014 and CA-015!")

if __name__ == '__main__':
    create_doc()
